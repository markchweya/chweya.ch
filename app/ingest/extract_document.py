"""Extraction for uploaded formats the crawler pipeline does not cover.

HTML and PDF already have extractors, used for both crawled and uploaded
content. This module adds DOCX, plain text, Markdown and CSV.

The same rule applies as everywhere else in ingestion: remove and preserve,
never rewrite. A fee written "CHF 20.--" stays exactly that.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field

from app.db.models.content import ExtractionQuality
from app.ingest.extract_html import ExtractedBlock, ExtractedPage
from app.observability import get_logger

logger = get_logger(__name__)

# Below this a document carries nothing worth indexing.
MIN_DOCUMENT_CHARACTERS = 120

# A CSV wider than this is a data export, not reference material a resident
# would read. Rendering a hundred columns as prose helps nobody.
MAX_CSV_COLUMNS = 20
MAX_CSV_ROWS = 2000

# Markdown heading, and the setext form some editors still emit.
ATX_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
SETEXT_UNDERLINE = re.compile(r"^(=+|-+)\s*$")


@dataclass
class ExtractedDocument:
    """Text and structure pulled from an uploaded file."""

    title: str = ""
    blocks: list[ExtractedBlock] = field(default_factory=list)
    quality: ExtractionQuality = ExtractionQuality.GOOD
    notes: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks)

    @property
    def is_usable(self) -> bool:
        return len(self.text) >= MIN_DOCUMENT_CHARACTERS

    def as_page(self) -> ExtractedPage:
        """Adapt to the shape the chunker already understands."""
        page = ExtractedPage(title=self.title, blocks=self.blocks)
        if not self.is_usable:
            page.quality_note = "insufficient_text"
        return page


def _decode(data: bytes) -> tuple[str, str]:
    """Decode bytes, reporting which encoding worked.

    Tried in order of likelihood for Swiss administrative documents. cp1252
    last, because it decodes almost anything and would mask a real problem if
    tried first.
    """
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
        try:
            return data.decode(encoding), encoding
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-with-replacement"


def extract_text(data: bytes, *, filename: str | None = None) -> ExtractedDocument:
    """Extract from a plain text file.

    Paragraphs are separated by blank lines, which is the only structure a
    .txt file carries. Nothing is inferred beyond that.
    """
    result = ExtractedDocument()
    content, encoding = _decode(data)
    if encoding == "utf-8-with-replacement":
        result.quality = ExtractionQuality.PARTIAL
        result.notes.append("text_contained_undecodable_bytes")

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", content) if part.strip()]
    for paragraph in paragraphs:
        collapsed = re.sub(r"[ \t]+", " ", paragraph.replace("\r\n", "\n")).strip()
        if collapsed:
            result.blocks.append(ExtractedBlock(text=collapsed, tag="p"))

    # A .txt has no title, so the first short line is the best guess, and the
    # filename is the fallback.
    if paragraphs and len(paragraphs[0]) < 100 and "\n" not in paragraphs[0]:
        result.title = paragraphs[0]
    elif filename:
        result.title = filename.rsplit(".", 1)[0].replace("_", " ").strip()

    if not result.is_usable:
        result.quality = ExtractionQuality.LOW
        result.notes.append("insufficient_text")
    return result


def extract_markdown(data: bytes, *, filename: str | None = None) -> ExtractedDocument:
    """Extract from Markdown, keeping the heading structure.

    Markdown carries real headings, so the section trail a citation needs can
    be built properly rather than guessed.
    """
    result = ExtractedDocument()
    content, _ = _decode(data)
    lines = content.replace("\r\n", "\n").split("\n")

    section_path: list[str] = []
    buffer: list[str] = []
    in_code_fence = False

    def flush() -> None:
        if not buffer:
            return
        text = " ".join(line.strip() for line in buffer).strip()
        buffer.clear()
        if text:
            result.blocks.append(
                ExtractedBlock(text=text, tag="p", section_path=tuple(section_path))
            )

    for index, raw in enumerate(lines):
        line = raw.rstrip()

        # Code fences are passed through untouched. Reflowing a code block
        # changes what it says.
        if line.strip().startswith("```"):
            in_code_fence = not in_code_fence
            continue
        if in_code_fence:
            buffer.append(raw)
            continue

        heading = ATX_HEADING.match(line)
        # A setext heading is the previous line underlined with = or -.
        if (
            not heading
            and SETEXT_UNDERLINE.match(line)
            and index > 0
            and lines[index - 1].strip()
            and buffer
        ):
            level = 1 if line.strip().startswith("=") else 2
            title = buffer.pop().strip()
            flush()
            del section_path[level - 1 :]
            while len(section_path) < level - 1:
                section_path.append("")
            section_path.append(title)
            result.blocks.append(
                ExtractedBlock(text=title, tag=f"h{level}", section_path=tuple(section_path))
            )
            if not result.title and level == 1:
                result.title = title
            continue

        if heading:
            flush()
            level = len(heading.group(1))
            title = heading.group(2).strip()
            del section_path[level - 1 :]
            while len(section_path) < level - 1:
                section_path.append("")
            section_path.append(title)
            result.blocks.append(
                ExtractedBlock(text=title, tag=f"h{level}", section_path=tuple(section_path))
            )
            if not result.title and level == 1:
                result.title = title
            continue

        if not line.strip():
            flush()
            continue

        # A list item is its own block, so a requirements list keeps its items
        # separate rather than becoming one paragraph.
        if re.match(r"^\s*([-*+]|\d+\.)\s+", line):
            flush()
            item = re.sub(r"^\s*([-*+]|\d+\.)\s+", "", line).strip()
            if item:
                result.blocks.append(
                    ExtractedBlock(text=item, tag="li", section_path=tuple(section_path))
                )
            continue

        buffer.append(line)

    flush()

    if not result.title and filename:
        result.title = filename.rsplit(".", 1)[0].replace("_", " ").strip()
    if not result.is_usable:
        result.quality = ExtractionQuality.LOW
        result.notes.append("insufficient_text")
    return result


def extract_csv(data: bytes, *, filename: str | None = None) -> ExtractedDocument:
    """Extract from CSV, rendering each row as a labelled line.

    Section 7 admits CSV only when it holds useful public reference
    information, such as a fee schedule. Each row becomes "Column: value"
    pairs, because a bare row of values is meaningless once retrieved out of
    the context of its header.
    """
    result = ExtractedDocument()
    content, _ = _decode(data)

    try:
        # Sniffing handles the semicolon delimiter, which is what a
        # German-locale Excel export produces and which a comma-only reader
        # turns into one column per row.
        dialect = csv.Sniffer().sniff(content[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel  # type: ignore[assignment]

    reader = csv.reader(io.StringIO(content), dialect)
    try:
        rows = list(reader)
    except csv.Error as exc:
        result.quality = ExtractionQuality.FAILED
        result.notes.append(f"csv_parse_failed: {type(exc).__name__}")
        return result

    rows = [row for row in rows if any(cell.strip() for cell in row)]
    if not rows:
        result.quality = ExtractionQuality.FAILED
        result.notes.append("csv_empty")
        return result

    header = [cell.strip() for cell in rows[0]]
    if len(header) > MAX_CSV_COLUMNS:
        result.quality = ExtractionQuality.LOW
        result.notes.append(f"csv_too_wide: {len(header)} columns")
        return result

    if len(rows) - 1 > MAX_CSV_ROWS:
        result.notes.append(f"csv_truncated_at_{MAX_CSV_ROWS}_rows")
        rows = rows[: MAX_CSV_ROWS + 1]

    result.blocks.append(ExtractedBlock(text=", ".join(header), tag="h1"))

    for row in rows[1:]:
        pairs = [
            f"{header[index]}: {cell.strip()}"
            for index, cell in enumerate(row)
            if index < len(header) and cell.strip()
        ]
        if pairs:
            result.blocks.append(
                ExtractedBlock(text="; ".join(pairs), tag="li", section_path=(", ".join(header),))
            )

    if filename:
        result.title = filename.rsplit(".", 1)[0].replace("_", " ").strip()
    if not result.is_usable:
        result.quality = ExtractionQuality.LOW
        result.notes.append("insufficient_text")
    return result


def extract_docx(data: bytes, *, filename: str | None = None) -> ExtractedDocument:
    """Extract from a Word document, keeping headings and tables.

    Validation has already confirmed the container holds a Word document and
    carries no macros or embedded executables. python-docx reads XML and
    executes nothing regardless.
    """
    result = ExtractedDocument()

    try:
        import docx
    except ImportError:  # pragma: no cover - the worker image installs it
        result.quality = ExtractionQuality.FAILED
        result.notes.append("python_docx_not_installed")
        return result

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - a malformed file is a finding
        result.quality = ExtractionQuality.FAILED
        result.notes.append(f"docx_unreadable: {type(exc).__name__}")
        return result

    core = document.core_properties
    result.title = (core.title or "").strip()

    section_path: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        heading = re.match(r"heading (\d)", style)

        if heading:
            level = int(heading.group(1))
            del section_path[level - 1 :]
            while len(section_path) < level - 1:
                section_path.append("")
            section_path.append(text)
            result.blocks.append(
                ExtractedBlock(text=text, tag=f"h{min(level, 6)}", section_path=tuple(section_path))
            )
            if not result.title and level == 1:
                result.title = text
            continue

        tag = "li" if "list" in style else "p"
        result.blocks.append(
            ExtractedBlock(text=text, tag=tag, section_path=tuple(section_path))
        )

    # Tables carry fee schedules and opening hours, which is exactly the
    # content residents ask about, so each row is kept with its header labels.
    for table in document.tables:
        rows = table.rows
        if not rows:
            continue
        header = [cell.text.strip() for cell in rows[0].cells]
        for row in rows[1:]:
            pairs = [
                f"{header[index]}: {cell.text.strip()}"
                for index, cell in enumerate(row.cells)
                if index < len(header) and cell.text.strip() and header[index]
            ]
            if pairs:
                result.blocks.append(
                    ExtractedBlock(
                        text="; ".join(pairs), tag="li", section_path=tuple(section_path)
                    )
                )

    if not result.title and filename:
        result.title = filename.rsplit(".", 1)[0].replace("_", " ").strip()

    if not result.blocks:
        result.quality = ExtractionQuality.FAILED
        result.notes.append("no_text_in_document")
    elif not result.is_usable:
        result.quality = ExtractionQuality.LOW
        result.notes.append("insufficient_text")

    return result
